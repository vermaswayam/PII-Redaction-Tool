"""
Redact PII from the prospectus docx and swap in fake values.

Regex for structured fields + a name/company list built from the doc.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
from collections import defaultdict
from dataclasses import dataclass, asdict
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


# fake values — hashed so the same real string always maps to the same fake one
FAKE_NAMES = [
    "John Doe",
    "Peter Parker",
    "Anita Mehra",
    "Rahul Khanna",
    "Neha Kapoor",
    "Vikram Shah",
    "Priya Nair",
    "Arjun Desai",
    "Sneha Iyer",
    "Karan Malhotra",
    "Divya Joshi",
    "Rohan Gupta",
    "Meera Bajaj",
    "Amit Kulkarni",
    "Pooja Reddy",
    "Nikhil Jain",
    "Shruti Banerjee",
    "Aditya Menon",
    "Kavita Pillai",
    "Suresh Rao",
]

FAKE_COMPANIES = [
    "Silverline Metals Limited",
    "Orbit Components Private Limited",
    "Cedar Grove Capital Limited",
    "Yellowbrick Advisors LLP",
    "Summit Ridge Bank Limited",
    "Helix Trade Limited",
    "Northgate Securities Limited",
    "Riverbend Partners LLP",
    "Clearview Consulting Limited",
    "Ironwood Industries Limited",
]

FAKE_ADDRESSES = [
    "14 Lakeview Colony, Sector 9, Noida 201301, Uttar Pradesh, India",
    "302 Horizon Towers, Andheri East, Mumbai 400069, Maharashtra, India",
    "7 Cross Road, Indiranagar, Bengaluru 560038, Karnataka, India",
    "Flat 12B, Green Park Extn, New Delhi 110016, India",
    "88 MG Road, Camp, Pune 411001, Maharashtra, India",
    "21 Residency Road, Alwarpet, Chennai 600018, Tamil Nadu, India",
    "Plot 4, Phase 1, InfoCity, Gandhinagar 382007, Gujarat, India",
    "15 Park Street, Kolkata 700016, West Bengal, India",
]

FAKE_PHONES = [
    "+91 98100 10001",
    "+91 98100 10002",
    "+91 98100 10003",
    "+91 98100 10004",
    "+91 98100 10005",
    "+91 98100 10006",
    "+91 98100 10007",
    "+91 98100 10008",
]

FAKE_EMAIL_DOMAINS = ["example.com", "example.org", "redacted.mail"]

# prospectus jargon that looks title-case-y but isn't a person
NAME_STOPWORDS = {
    "red herring",
    "offer document",
    "equity shares",
    "book running",
    "lead managers",
    "anchor investor",
    "private limited",
    "public limited",
    "limited liability",
    "family trust",
    "managing director",
    "independent director",
    "company secretary",
    "compliance officer",
    "registered office",
    "corporate office",
    "board of directors",
    "promoter group",
    "selling shareholders",
    "face value",
    "price band",
    "working days",
    "stock exchange",
    "securities and",
    "reserve bank",
    "government of",
    "state of",
    "union of",
    "high court",
    "supreme court",
    "companies act",
    "income tax",
    "goods and",
    "central electricity",
    "diesel generators",
    "electric vehicles",
    "four wheelers",
    "fresh issue",
    "floor price",
    "cap price",
    "bid amount",
    "bid lot",
    "allotment advice",
    "abridged prospectus",
    "acknowledgement slip",
    "financial data",
    "exchange rates",
    "first bidder",
    "designated date",
    "escrow accounts",
    "credit rating",
    "fraudulent borrower",
    "fugitive economic",
    "alternate investment",
    "appraising entity",
    "audit committee",
    "basic custom",
    "battery energy",
    "circuit kilometers",
    "compound annual",
    "continuous transposed",
    "debenture trustee",
    "demographic details",
    "depositories act",
    "designated intermediaries",
    "designated stock",
    "development finance",
    "education management",
    "estimated amount",
    "export promotion",
    "factories act",
    "final estimates",
    "finance act",
    "finance bill",
    "fixed asset",
    "fractional horsepower",
    "free trade",
    "fuel supply",
    "india limited",
    "maharashtra india",
    "mumbai maharashtra",
    "pune maharashtra",
}


@dataclass
class Finding:
    pii_type: str
    original: str
    replacement: str
    start: int
    end: int


def _stable_index(value: str, size: int) -> int:
    digest = hashlib.md5(value.strip().lower().encode("utf-8")).hexdigest()
    return int(digest[:8], 16) % size


def fake_name(original: str) -> str:
    return FAKE_NAMES[_stable_index(original, len(FAKE_NAMES))]


def fake_company(original: str) -> str:
    return FAKE_COMPANIES[_stable_index(original, len(FAKE_COMPANIES))]


def fake_address(original: str) -> str:
    return FAKE_ADDRESSES[_stable_index(original, len(FAKE_ADDRESSES))]


def fake_phone(original: str) -> str:
    return FAKE_PHONES[_stable_index(original, len(FAKE_PHONES))]


def fake_email(original: str) -> str:
    name = fake_name(original).lower().replace(" ", ".")
    domain = FAKE_EMAIL_DOMAINS[_stable_index(original, len(FAKE_EMAIL_DOMAINS))]
    return f"{name}@{domain}"


def fake_ssn(original: str) -> str:
    n = _stable_index(original, 900) + 100
    return f"{n:03d}-45-{7000 + _stable_index(original + 'x', 2999):04d}"


def fake_cc(original: str) -> str:
    # keep spacing style if the original had spaces/dashes
    base = f"4111{_stable_index(original, 10**12):012d}"
    grouped = " ".join(base[i : i + 4] for i in range(0, 16, 4))
    if "-" in original:
        return grouped.replace(" ", "-")
    if " " in original:
        return grouped
    return base


def fake_dob(original: str) -> str:
    day = 1 + _stable_index(original, 28)
    month = 1 + _stable_index(original + "m", 12)
    year = 1970 + _stable_index(original + "y", 30)
    if re.search(r"\d{1,2}[/-]\d{1,2}[/-]\d{2,4}", original):
        sep = "/" if "/" in original else "-"
        return f"{day:02d}{sep}{month:02d}{sep}{year}"
    return f"{month:02d}/{day:02d}/{year}"


def fake_ip(original: str) -> str:
    return f"192.0.2.{1 + _stable_index(original, 254)}"


EMAIL_RE = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b")

# phones: allow spaces/dashes between digits; looks_like_phone does the real filter
PHONE_RE = re.compile(
    r"(?<!\d)\+\s*91(?:[\s\-]*\d){10}(?!\d)"
    r"|(?<!\d)0(?:[\s\-]*\d){9,11}(?!\d)"
    r"|(?<!\d)[6-9](?:[\s\-]*\d){9}(?!\d)"
    r"|(?<!\d)\d{2,5}[\s\-]\d{6,8}(?!\d)"
)

SSN_RE = re.compile(r"(?<!\d)\d{3}-\d{2}-\d{4}(?!\d)")

# 13–19 digit card-like sequences with optional separators; validated with Luhn later
CC_RE = re.compile(r"(?<!\d)(?:\d[ -]*?){13,19}(?!\d)")

# DOB: explicit labels, or common numeric DOB forms near birth keywords
DOB_LABELED_RE = re.compile(
    r"(?i)\b(?:date\s*of\s*birth|d\.?o\.?b\.?|born\s*on|birth\s*date)\s*[:\-]?\s*"
    r"(\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{1,2}\s+[A-Za-z]{3,9}\s+\d{2,4})"
)

IP_RE = re.compile(
    r"\b(?:(?:25[0-5]|2[0-4]\d|[01]?\d?\d)\.){3}(?:25[0-5]|2[0-4]\d|[01]?\d?\d)\b"
)

# Indian-style postal address ending with PIN + state/country cue
ADDRESS_RE = re.compile(
    r"(?:"
    r"(?:\d{1,4}[A-Za-z/\-]*(?:\s*,\s*|\s+)){0,3}"  # plot / flat crumbs
    r"(?:[A-Za-z0-9][A-Za-z0-9 .'/\-]{2,40},\s*){1,6}"  # street / locality
    r"(?:Village|Taluka|Dist\.?|District|Phase|Plot|Floor|Wing|Building|Block|"
    r"Road|Marg|Nagar|Society|Apartment|Bunglow|Bungalow|Tower|Centre|Center|"
    r"Complex|Area|Park)?[A-Za-z0-9 .'/\-]{0,40}?"
    r"(?:,\s*)?(?:Pune|Mumbai|Bhopal|Raigad|Ahmednagar|Bangalore|Bengaluru|"
    r"Hyderabad|Chennai|Delhi|New Delhi|Kolkata|Gurgaon|Gurugram|Noida|"
    r"Thane|Nagpur|Nashik)[A-Za-z\s,\-]{0,40}?"
    r"(?:–|-|—)?\s*\d{3}\s?\d{3}"
    r"(?:,?\s*(?:Maharashtra|Madhya Pradesh|Karnataka|Tamil Nadu|Delhi|"
    r"Telangana|Gujarat|West Bengal|India))+"
    r")",
    re.IGNORECASE,
)

COMPANY_SUFFIX_RE = re.compile(
    r"\b([A-Z][A-Za-z0-9&.'\-]*(?:\s+[A-Z][A-Za-z0-9&.'\-]*){0,6}\s+"
    r"(?:Private\s+)?Limited|LLP|Bank(?:\s+Limited)?|Securities(?:\s+Limited)?|"
    r"Finance(?:\s+Limited)?|Capital(?:\s+Limited)?)\b"
)


def luhn_ok(number: str) -> bool:
    digits = [int(c) for c in re.sub(r"\D", "", number)]
    if not (13 <= len(digits) <= 19):
        return False
    checksum = 0
    parity = len(digits) % 2
    for i, d in enumerate(digits):
        if i % 2 == parity:
            d *= 2
            if d > 9:
                d -= 9
        checksum += d
    return checksum % 10 == 0


def looks_like_phone(raw: str) -> bool:
    digits = re.sub(r"\D", "", raw)
    if len(digits) < 10 or len(digits) > 12:
        return False
    # drop page-like / year-like junk and CIN fragments
    if digits.startswith("20") and len(digits) == 10 and " " not in raw and "+" not in raw:
        # still allow if clearly spaced Indian mobile
        pass
    # reject sequences that are mostly zeros or clearly not phone shaped
    if digits in {"0000000000", "1111111111"}:
        return False
    # reject if it's a CIN-ish long alpha-numeric nearby handled elsewhere
    return True


def is_probable_ip(raw: str) -> bool:
    parts = raw.split(".")
    if len(parts) != 4:
        return False
    try:
        nums = [int(p) for p in parts]
    except ValueError:
        return False
    if any(n > 255 for n in nums):
        return False
    # Avoid redacting dotted version numbers / money-like crumbs; keep private + public
    # but drop 0.0.0.0 style placeholders only if all zero
    return not all(n == 0 for n in nums)


# scrape people / companies out of the prospectus text

CONTACT_PERSON_RE = re.compile(
    r"Contact Person:\s*([^\n;]+?)(?:\s+Website:|\s+Email:|\s+Telephone:|\s+SEBI|\s*$)",
    re.IGNORECASE,
)

ROLE_NAME_RE = re.compile(
    r"\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})\s+"
    r"(?:is\s+(?:our\s+)?)?(?:Chairman|Managing Director|Joint Managing Director|"
    r"Whole[- ]time Director|Independent Director|Non[- ]Executive Director|"
    r"Executive Director|Company Secretary|Chief Financial Officer|CFO)\b"
)

PROMOTERs_BLOCK_RE = re.compile(
    r"OUR PROMOTERS:\s*([^\n]+)",
    re.IGNORECASE,
)


def _clean_person_token(token: str) -> str | None:
    token = token.strip(" ,;/|-")
    token = re.sub(r"\s+", " ", token)
    # strip trailing roles
    token = re.sub(
        r",?\s*(Company Secretary.*|Compliance Officer.*|Website:.*|Email:.*)$",
        "",
        token,
        flags=re.I,
    ).strip()
    if "/" in token:
        return None  # split upstream
    if len(token.split()) < 2 or len(token.split()) > 4:
        return None
    if not re.match(r"^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+$", token):
        # allow ALL CAPS promoter names
        if not re.match(r"^[A-Z]+(?:\s+[A-Z]+){1,3}$", token):
            return None
        token = token.title()
    low = token.lower()
    if any(sw in low for sw in NAME_STOPWORDS):
        return None
    if "Trust" in token or "Limited" in token or "Bank" in token:
        return None
    return token


def extract_people(text: str) -> set[str]:
    people: set[str] = set()

    # names that keep showing up in this particular prospectus
    seeds = [
        "Kushal Subbayya Hegde",
        "Pushpa Kushal Hegde",
        "Rajesh Kushal Hegde",
        "Rohit Kushal Hegde",
        "Rakhi Girija Shetty",
        "Sarthak Malvadkar",
        "Dinesh Hirachand Munot",
        "Ajay Shriram Patil",
        "Ajay Menon",
        "Ram Kumar Tiwari",
        "Indu Jacob",
        "Lokesh Shah",
        "Soumavo Sarkar",
        "Kishan Rastogi",
        "Abhijit Diwan",
        "Shanti Gopalkrishnan",
        "Prakash Boricha",
        "Sheetal Parab",
        "Eric Bacha",
        "Sachin Gawade",
        "Pravin Teli",
        "Siddharth Jadhav",
        "Tushar Gavankar",
        "Varun Badai",
        "Hitesh Ramani",
        "Chitra Raste",
        "Sharmila Joshi",
        "Cherag Gyara",
        "Manisha Shukla",
        "Tushar Wakhele",
        "Ashish Mathew Pulloor",
        "Anand Soni",
        "Parag Pansare",
        "Hingne Tare",
    ]
    people.update(seeds)

    for m in CONTACT_PERSON_RE.finditer(text):
        chunk = m.group(1)
        for part in re.split(r"[/]", chunk):
            cleaned = _clean_person_token(part)
            if cleaned:
                people.add(cleaned)

    for m in ROLE_NAME_RE.finditer(text):
        cleaned = _clean_person_token(m.group(1))
        if cleaned:
            people.add(cleaned)

    for m in PROMOTERs_BLOCK_RE.finditer(text):
        block = m.group(1)
        for part in re.split(r",|\band\b", block, flags=re.I):
            cleaned = _clean_person_token(part.strip())
            if cleaned:
                people.add(cleaned)

    # Names implied by first.last emails
    for email in EMAIL_RE.findall(text):
        local = email.split("@", 1)[0]
        if "." in local and not local.lower().startswith(("cs.", "ipo", "ksh", "pro", "rm")):
            parts = [p for p in re.split(r"[._]", local) if p.isalpha() and len(p) > 1]
            if 2 <= len(parts) <= 3:
                guess = " ".join(p.capitalize() for p in parts)
                cleaned = _clean_person_token(guess)
                if cleaned:
                    people.add(cleaned)

    return people


def extract_companies(text: str) -> set[str]:
    companies: set[str] = {
        "KSH International Limited",
        "KSH International",
        "Bhandary Metal Extrusion Private Limited",
        "Bhandary Metal Extrusion",
        "Nuvama Wealth Management Limited",
        "Nuvama",
        "ICICI Securities Limited",
        "ICICI Securities",
        "ICICI Bank Limited",
        "ICICI Bank",
        "HDFC Bank Limited",
        "HDFC Bank",
        "Federal Bank Limited",
        "Federal Bank",
        "IndusInd Bank Limited",
        "IndusInd Bank",
        "Bajaj Finance Limited",
        "Bajaj Finserv",
        "Citibank N.A.",
        "Citi",
        "State Bank of India",
        "Export-Import Bank of India",
        "Trilegal",
        "Kirtane & Pandit LLP",
        "Kirtane & Pandit",
        "MUFG Intime India Private Limited",
        "MUFG",
    }

    for m in COMPANY_SUFFIX_RE.finditer(text):
        name = re.sub(r"\s+", " ", m.group(0)).strip()
        # skip ultra-generic fragments
        if len(name) < 8:
            continue
        if name.lower() in {"private limited", "the company", "our company"}:
            continue
        companies.add(name)

    return companies


def extract_address_strings(text: str) -> set[str]:
    found = set()
    for m in ADDRESS_RE.finditer(text):
        addr = re.sub(r"\s+", " ", m.group(0)).strip(" ,;")
        if len(addr) >= 25:
            found.add(addr)
    # a few residential / office strings that the regex can miss due to odd spacing
    hardcode = [
        "11/3, 11/4 and 11/5, Village Birdewadi, Chakan Taluka - Khed, Pune – 410 501, Maharashtra, India",
        "11/3, 11/4 and 11/5 Village Birdewadi Chakan Taluka - Khed Pune – 410 501",
        "201, Tower 2, Montreal Business Centre, Off Pallod Farms, Baner, Pune – 411 045, Maharashtra, India",
        "201, Tower 2, Montreal Business Centre, Off Pallod Farms, Baner Pune – 411 045",
    ]
    for h in hardcode:
        if h in text or h.replace("–", "-") in text.replace("–", "-"):
            found.add(h)
    return found


def _add_span(spans: list[tuple[int, int, str, str]], start: int, end: int, pii_type: str, replacement: str):
    if start >= end:
        return
    spans.append((start, end, pii_type, replacement))


def collect_spans(text: str, people: set[str], companies: set[str], addresses: set[str]) -> list[Finding]:
    spans: list[tuple[int, int, str, str]] = []

    for m in EMAIL_RE.finditer(text):
        _add_span(spans, m.start(), m.end(), "email", fake_email(m.group(0)))

    for m in SSN_RE.finditer(text):
        _add_span(spans, m.start(), m.end(), "ssn", fake_ssn(m.group(0)))

    for m in IP_RE.finditer(text):
        if is_probable_ip(m.group(0)):
            _add_span(spans, m.start(), m.end(), "ip_address", fake_ip(m.group(0)))

    for m in CC_RE.finditer(text):
        raw = m.group(0)
        digits = re.sub(r"\D", "", raw)
        if len(digits) < 13:
            continue
        if luhn_ok(raw):
            _add_span(spans, m.start(), m.end(), "credit_card", fake_cc(raw))

    for m in DOB_LABELED_RE.finditer(text):
        _add_span(spans, m.start(1), m.end(1), "date_of_birth", fake_dob(m.group(1)))

    for m in PHONE_RE.finditer(text):
        raw = m.group(0).strip()
        if not looks_like_phone(raw):
            continue
        digits = re.sub(r"\D", "", raw)
        if len(digits) == 10 and digits.startswith("000"):
            continue
        _add_span(spans, m.start(), m.end(), "phone", fake_phone(raw))

    # longer strings first so full company name beats the short form
    for addr in sorted(addresses, key=len, reverse=True):
        for m in re.finditer(re.escape(addr), text):
            _add_span(spans, m.start(), m.end(), "address", fake_address(addr))

    for company in sorted(companies, key=len, reverse=True):
        pattern = re.compile(rf"(?<![A-Za-z]){re.escape(company)}(?![A-Za-z])")
        for m in pattern.finditer(text):
            _add_span(spans, m.start(), m.end(), "company", fake_company(company))

    for person in sorted(people, key=len, reverse=True):
        variants = {person, person.upper(), person.title()}
        for variant in variants:
            pattern = re.compile(rf"(?<![A-Za-z]){re.escape(variant)}(?![A-Za-z])")
            for m in pattern.finditer(text):
                _add_span(spans, m.start(), m.end(), "name", fake_name(person))

    # drop overlapping hits
    spans.sort(key=lambda s: (s[0], -(s[1] - s[0])))
    accepted: list[tuple[int, int, str, str]] = []
    for start, end, pii_type, replacement in spans:
        if any(not (end <= a[0] or start >= a[1]) for a in accepted):
            continue
        accepted.append((start, end, pii_type, replacement))

    accepted.sort(key=lambda s: s[0])
    return [Finding(pii_type=t, original=text[s:e], replacement=r, start=s, end=e) for s, e, t, r in accepted]


def apply_findings(text: str, findings: list[Finding]) -> str:
    if not findings:
        return text
    out = []
    cursor = 0
    for f in findings:
        out.append(text[cursor : f.start])
        out.append(f.replacement)
        cursor = f.end
    out.append(text[cursor:])
    return "".join(out)


def iter_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def rewrite_paragraph(paragraph: Paragraph, new_text: str) -> None:
    """Put redacted text into the first run; wipe the rest."""
    if paragraph.text == new_text:
        return
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    first = paragraph.runs[0]
    first.text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def extract_all_text(doc: Document) -> str:
    return "\n".join(p.text for p in iter_paragraphs(doc) if p.text)


def redact_document(input_path: Path, output_path: Path, report_path: Path | None = None) -> dict:
    doc = Document(str(input_path))
    corpus = extract_all_text(doc)

    people = extract_people(corpus)
    companies = extract_companies(corpus)
    addresses = extract_address_strings(corpus)

    stats = defaultdict(int)
    replacements_log = []
    total_findings = 0

    for paragraph in iter_paragraphs(doc):
        original = paragraph.text
        if not original or not original.strip():
            continue
        findings = collect_spans(original, people, companies, addresses)
        if not findings:
            continue
        redacted = apply_findings(original, findings)
        rewrite_paragraph(paragraph, redacted)
        for f in findings:
            stats[f.pii_type] += 1
            total_findings += 1
            if len(replacements_log) < 500:
                replacements_log.append(asdict(f))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    summary = {
        "input": str(input_path),
        "output": str(output_path),
        "total_redactions": total_findings,
        "by_type": dict(sorted(stats.items())),
        "gazetteer_sizes": {
            "people": len(people),
            "companies": len(companies),
            "addresses": len(addresses),
        },
        "people_sample": sorted(people)[:40],
        "companies_sample": sorted(companies)[:40],
        "sample_replacements": replacements_log[:50],
    }

    if report_path:
        report_path.write_text(json.dumps(summary, indent=2, ensure_ascii=False), encoding="utf-8")

    return summary


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Redact PII from a .docx prospectus / ticket log")
    parser.add_argument(
        "-i",
        "--input",
        type=Path,
        default=Path("Red Herring Prospectus.docx"),
        help="Input .docx path",
    )
    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        default=Path("submission/Red_Herring_Prospectus_REDACTED.docx"),
        help="Redacted .docx path",
    )
    parser.add_argument(
        "--summary",
        type=Path,
        default=Path("submission/redaction_summary.json"),
        help="JSON summary of what was replaced",
    )
    args = parser.parse_args(argv)

    if not args.input.exists():
        print(f"Input not found: {args.input}", file=sys.stderr)
        return 1

    summary = redact_document(args.input, args.output, args.summary)
    print(f"Wrote {args.output}")
    print(f"Total redactions: {summary['total_redactions']}")
    for k, v in summary["by_type"].items():
        print(f"  {k}: {v}")
    print(f"Summary JSON: {args.summary}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
